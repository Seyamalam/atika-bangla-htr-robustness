from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
OUT = ROOT / "Bangla_HTR_Robustness_Manuscript.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_run(run, bold=False, size=11, color="000000") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        style_run(run, bold=True, size=16 if level == 1 else 14 if level == 2 else 12)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    style_run(run)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    style_run(run, size=9, color="555555")


def main() -> None:
    summary = json.loads((RESULTS / "summary.json").read_text(encoding="utf-8"))
    hf_status = json.loads((RESULTS / "hf_access_status.json").read_text(encoding="utf-8"))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Cross-Dataset Robustness of Bangla Handwritten Text Recognition")
    style_run(run, bold=True, size=22)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Reproducible preliminary manuscript and dataset-analysis report")
    style_run(run, size=12, color="555555")

    add_heading(doc, "Abstract", 1)
    doc.add_paragraph(
        "This manuscript documents the reproducible setup for a thesis on Bangla handwritten text "
        "recognition robustness. The current stage verifies BN-HTRd dataset archives, profiles sample "
        "line/word images, runs preprocessing diagnostics, and defines the controlled evaluation path "
        "needed to measure cross-dataset performance drop."
    )

    add_heading(doc, "Research Question", 1)
    doc.add_paragraph(
        "How robust are Bangla HTR models trained on BN-HTRd when evaluated outside their original "
        "benchmark distribution?"
    )
    add_bullet(doc, "Measure the gap between benchmark and real-world handwriting.")
    add_bullet(doc, "Keep the external dataset separate from training and tuning.")
    add_bullet(doc, "Analyze failure modes by skew, noise, lighting, stroke density, and segmentation quality.")

    add_heading(doc, "Methods", 1)
    steps = [
        "Verify all downloaded dataset archives with SHA-256.",
        "Build archive and sample-level inventories.",
        "Profile sample image quality with width, height, contrast, blur proxy, and ink fraction.",
        "Run preprocessing diagnostics using denoising, CLAHE, and adaptive thresholding.",
        "Prepare writer-safe split planning and an external annotation protocol.",
    ]
    for step in steps:
        add_bullet(doc, step)

    add_heading(doc, "Dataset Inventory", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Archive", "Members", "JPG", "TXT", "XML"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "F2F4F7")
        r = cell.paragraphs[0].add_run(h)
        style_run(r, bold=True)
    import csv

    with (RESULTS / "archive_inventory.csv").open(newline="", encoding="utf-8") as fh:
        for row in csv.DictReader(fh):
            cells = table.add_row().cells
            values = [row["archive"], row["members"], row["jpg"], row["txt"], row["xml"]]
            for cell, value in zip(cells, values):
                r = cell.paragraphs[0].add_run(str(value))
                style_run(r, size=9)

    add_heading(doc, "Computed Preliminary Results", 1)
    add_bullet(doc, f"Sample files extracted/profiled: {summary['sample_files']:,}.")
    add_bullet(doc, f"Sample JPEG images: {summary['sample_jpg']:,}.")
    add_bullet(doc, f"Ground-truth words in sample documents: {summary['sample_text_words']:,}.")
    add_bullet(doc, f"Line/word images profiled: {summary['line_images_profiled']:,}.")
    add_bullet(doc, f"Images preprocessed for comparison: {summary['preprocessed_images']:,}.")

    for image_name, caption in [
        ("sample_image_distributions.png", "Figure 1. Sample image dimensions and estimated ink-density distributions."),
        ("preprocessing_ink_shift.png", "Figure 2. Ink-density shift after preprocessing."),
    ]:
        image_path = RESULTS / image_name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, caption)

    add_heading(doc, "Access and Reproducibility Notes", 1)
    doc.add_paragraph(
        f"Hugging Face split status: {hf_status['status']}. Reason: {hf_status['reason']} "
        f"Next action: {hf_status['next_action']}"
    )
    doc.add_paragraph(
        "The final CER/WER phase requires either authorized access to the Hugging Face split or a local "
        "full extraction plus line-level label conversion. This manuscript records computed preliminary "
        "results only and does not claim final OCR accuracy."
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

